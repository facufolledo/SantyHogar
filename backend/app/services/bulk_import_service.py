"""Servicio para importaci├│n masiva de productos desde documentos."""
import re
import logging
from io import BytesIO
from typing import Dict, List, Optional, Tuple
from uuid import UUID

from docx import Document
from openpyxl import load_workbook
import zipfile

from app.models.bulk_import import (
    BulkImportResponse,
    ExcelImportConfirmRow,
    ExcelImportPreview,
    ProductImportRow,
    ProductImportValidation,
)

logger = logging.getLogger(__name__)


def _collect_docx_lines(doc: Document) -> list[str]:
    """P├írrafos + celdas de tablas (muchas planillas Word van en tablas, no en p├írrafos)."""
    lines: list[str] = []
    for p in doc.paragraphs:
        t = " ".join((p.text or "").split())
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = " ".join((cell.text or "").split())
                if t:
                    lines.append(t)
    return lines


def _extract_brand_from_name(nombre: str) -> str:
    """
    Extrae la marca del nombre del producto con tolerancia a errores tipogr├íficos.
    
    Ejemplos:
    - "HELADERA DREAN RDA250FVRT" ÔåÆ "Drean"
    - "EXHIBIDORA BRIKET 4300" ÔåÆ "Briket"
    - "EXHIBIDORA BRKET 4300" ÔåÆ "Briket" (detecta error tipogr├ífico)
    - "FREEZER HORIZONTAL INELRO FIH-270" ÔåÆ "Inelro"
    - "HELADERA VONDOM COMBI 250BI" ÔåÆ "Vondom"
    
    Args:
        nombre: Nombre completo del producto
        
    Returns:
        Marca extra├¡da o "Sin marca" si no se encuentra
    """
    # Lista de marcas conocidas (en may├║sculas) con variaciones comunes
    KNOWN_BRANDS = {
        'DREAN': ['DREAN', 'DEAN', 'DREN'],
        'BRIKET': ['BRIKET', 'BRKET', 'BRIKETT', 'BRIKT'],
        'INELRO': ['INELRO', 'INERL', 'INELR'],
        'VONDOM': ['VONDOM', 'VONDM', 'VONDON'],
        'KOHINOOR': ['KOHINOOR', 'KOHINOR', 'KHINOOR'],
        'NEBA': ['NEBA', 'NEEBA'],
        'SAMSUNG': ['SAMSUNG', 'SAMSUN', 'SAMUNG'],
        'LG': ['LG'],
        'WHIRLPOOL': ['WHIRLPOOL', 'WIRLPOOL', 'WHIRPOOL'],
        'PHILIPS': ['PHILIPS', 'PHILIPS', 'PHILLIPS'],
        'SONY': ['SONY'],
        'PANASONIC': ['PANASONIC', 'PANASONC'],
        'ELECTROLUX': ['ELECTROLUX', 'ELECTROLX'],
        'BOSCH': ['BOSCH', 'BOSH'],
        'SIEMENS': ['SIEMENS', 'SIEMEN'],
        'MABE': ['MABE'],
        'PATRICK': ['PATRICK', 'PATRIK'],
        'ATMA': ['ATMA'],
        'PHILCO': ['PHILCO', 'FILCO'],
        'NOBLEX': ['NOBLEX', 'NOBLX'],
        'BGH': ['BGH'],
        'SANYO': ['SANYO', 'SANIO'],
        'HITACHI': ['HITACHI', 'HITACH'],
        'TOSHIBA': ['TOSHIBA', 'TOSHBA'],
        'SHARP': ['SHARP', 'SHRP'],
        'DAEWOO': ['DAEWOO', 'DAEWU'],
        'AURORA': ['AURORA', 'AURRA'],
        'ESLAB├ôN': ['ESLAB├ôN', 'ESLABON', 'ESLABN'],
        'ESLABON': ['ESLABON', 'ESLAB├ôN', 'ESLABN'],
        'GAFA': ['GAFA'],
        'LONGVIE': ['LONGVIE', 'LONGVI'],
        'PEABODY': ['PEABODY', 'PEABDY'],
        'JAMES': ['JAMES', 'JMES'],
        'LILIANA': ['LILIANA', 'LILIAN'],
        'OSTER': ['OSTER', 'OSTR'],
        'MOULINEX': ['MOULINEX', 'MOULINX'],
        'TEFAL': ['TEFAL', 'TEFL'],
        'BLACK': ['BLACK', 'BLCK'],
        'DECKER': ['DECKER', 'DECKR'],
        'BRAUN': ['BRAUN', 'BRAU'],
        'ROWENTA': ['ROWENTA', 'ROWENT'],
        'KRUPS': ['KRUPS', 'KRUP'],
    }
    
    # Convertir a may├║sculas para comparaci├│n
    nombre_upper = nombre.upper()
    words = nombre_upper.split()
    
    # Buscar marca conocida en las palabras (incluyendo variaciones)
    for word in words:
        # Limpiar la palabra de caracteres especiales
        clean_word = re.sub(r'[^A-Z]', '', word)
        
        # Buscar coincidencia exacta o variaci├│n
        for brand_name, variations in KNOWN_BRANDS.items():
            if clean_word in variations:
                return brand_name.capitalize()
    
    # Si no se encuentra marca conocida, intentar detectar patr├│n
    # Generalmente la marca est├í despu├®s del tipo de producto
    # Ejemplo: "HELADERA DREAN RDA250" ÔåÆ segunda palabra
    product_types = ['HELADERA', 'FREEZER', 'LAVARROPAS', 'SECARROPAS', 
                     'LAVARROPA', 'EXHIBIDORA', 'FRIGOBAR', 'FRIOBAR',
                     'TV', 'TELEVISOR', 'SMART', 'LED', 'MICROONDAS',
                     'HORNO', 'COCINA', 'ANAFE', 'CAMPANA', 'PURIFICADOR',
                     'AIRE', 'ACONDICIONADO', 'VENTILADOR', 'ESTUFA',
                     'CALOVENTOR', 'TERMOTANQUE', 'CALEFON', 'CALEF├ôN']
    
    # Buscar tipo de producto y tomar la siguiente palabra como marca
    for i, word in enumerate(words):
        clean_word = re.sub(r'[^A-Z]', '', word)
        if clean_word in product_types and i + 1 < len(words):
            # La siguiente palabra podr├¡a ser la marca
            next_word = re.sub(r'[^A-Z]', '', words[i + 1])
            # Verificar que no sea un modelo (n├║meros o muy corto)
            if len(next_word) >= 3 and not next_word.isdigit():
                # Verificar si es una variaci├│n de marca conocida
                for brand_name, variations in KNOWN_BRANDS.items():
                    if next_word in variations:
                        return brand_name.capitalize()
                # Si no es conocida, retornar capitalizada
                return next_word.capitalize()
    
    # Si no se pudo detectar, retornar "Sin marca"
    return "Sin marca"


def generate_slug(nombre: str) -> str:
    """Genera un slug a partir del nombre del producto."""
    slug = nombre.lower()
    slug = re.sub(r'[^a-z0-9\s-]', '', slug)
    slug = re.sub(r'\s+', '-', slug)
    slug = re.sub(r'-+', '-', slug)
    return slug.strip('-')


# ------------------------------------------------------------------ #
# Mapeo de columnas para detecci├│n autom├ítica de headers Excel
# ------------------------------------------------------------------ #

COLUMN_ALIASES: Dict[str, List[str]] = {
    "nombre": ["nombre", "name", "producto", "product", "descripcion_producto", "articulo"],
    "categoria": ["categor├¡a", "categoria", "category", "cat"],
    "subcategoria": ["subcategor├¡a", "subcategoria", "subcategory", "sub_categoria", "sub"],
    "precio": ["precio", "price", "valor", "monto", "precio_venta"],
    "stock": ["stock", "cantidad", "qty", "quantity", "existencia", "existencias", "unidades"],
    "marca": ["marca", "brand", "fabricante", "manufacturer"],
    "descripcion": ["descripci├│n", "descripcion", "description", "detalle", "detalles", "desc"],
}


def _normalize_header(header: str) -> str:
    """Normaliza un header para comparaci├│n: lowercase, sin acentos comunes, sin espacios extra."""
    h = header.strip().lower()
    # Remover caracteres especiales excepto letras y n├║meros
    h = re.sub(r'[^a-z├í├®├¡├│├║├▒├╝0-9_]', '_', h)
    h = re.sub(r'_+', '_', h)
    return h.strip('_')


def _detect_column_mapping(headers: List[str]) -> Dict[str, int]:
    """
    Detecta autom├íticamente qu├® columna corresponde a qu├® campo de producto.
    
    Args:
        headers: Lista de nombres de columnas del Excel
        
    Returns:
        Diccionario {campo_producto: ├¡ndice_columna}
    """
    mapping: Dict[str, int] = {}
    normalized_headers = [_normalize_header(h) for h in headers]
    
    for field, aliases in COLUMN_ALIASES.items():
        for col_idx, norm_header in enumerate(normalized_headers):
            if norm_header in aliases:
                mapping[field] = col_idx
                break
            # Tambi├®n buscar coincidencia parcial
            for alias in aliases:
                if alias in norm_header or norm_header in alias:
                    if field not in mapping:
                        mapping[field] = col_idx
                    break
    
    return mapping


def parse_xlsx_file(file_content: bytes) -> List[ProductImportValidation]:
    """
    Parsea un archivo .xlsx usando openpyxl, detecta headers autom├íticamente
    y mapea columnas a campos de producto.
    
    Soporta dos formatos:
    1. Formato est├índar con headers en primera fila
    2. Formato de listado (sin headers): CODIGO - NOMBRE | CATEGORIA | STOCK
    
    Args:
        file_content: Contenido del archivo .xlsx en bytes
        
    Returns:
        Lista de validaciones de productos parseados
    """
    validations: List[ProductImportValidation] = []
    
    try:
        wb = load_workbook(BytesIO(file_content), read_only=True, data_only=True)
        ws = wb.active
        
        if ws is None:
            return [ProductImportValidation(
                row_number=0,
                valid=False,
                errors=["El archivo no contiene hojas de c├ílculo"]
            )]
        
        rows = list(ws.iter_rows(values_only=True))
        
        if len(rows) < 1:
            return [ProductImportValidation(
                row_number=0,
                valid=False,
                errors=["El archivo no contiene productos para importar"]
            )]
        
        # Detectar si es formato de listado (sin headers) o formato est├índar
        first_row = [str(cell) if cell is not None else "" for cell in rows[0]]
        first_row_text = " ".join(first_row).lower()
        
        # Si la primera fila tiene "Art├¡culo" o "Rubro", es un listado sin headers reales
        # O si tiene nombres propios como "Alejandro" (t├¡pico de reportes)
        is_listing_format = any(
            keyword in first_row_text
            for keyword in ["art├¡culo", "articulo", "rubro", "existencia", "listado", "alejandro", "fecha", "hora"]
        )
        
        # Tambi├®n detectar si NO hay headers t├¡picos de Excel (nombre, precio, stock, etc.)
        has_standard_headers = any(
            keyword in first_row_text
            for keyword in ["nombre", "name", "precio", "price", "stock", "cantidad", "categoria", "category"]
        )
        
        # Si no tiene headers est├índar, es formato de listado
        if not has_standard_headers:
            is_listing_format = True
        
        if is_listing_format:
            logger.info("Detectado formato de listado (sin headers est├índar)")
            validations = _parse_listing_format(rows)
        else:
            logger.info("Detectado formato est├índar con headers")
            validations = _parse_standard_format(rows)
        
        wb.close()
        
    except zipfile.BadZipFile:
        logger.error("El archivo no es un .xlsx v├ílido (no es un archivo ZIP)")
        validations.append(ProductImportValidation(
            row_number=0,
            valid=False,
            errors=["El archivo debe ser .xlsx (Excel 2007 o superior). Si es .xls, convertilo a .xlsx primero."]
        ))
    except Exception as e:
        logger.error(f"Error al parsear archivo Excel: {str(e)}")
        validations.append(ProductImportValidation(
            row_number=0,
            valid=False,
            errors=[f"Error al leer el archivo Excel: {str(e)}"]
        ))
    
    return validations


def _parse_listing_format(rows: list) -> List[ProductImportValidation]:
    """
    Parsea formato de listado sin headers est├índar.
    Formato esperado: CODIGO - NOMBRE | CATEGORIA - XX | STOCK
    """
    validations: List[ProductImportValidation] = []
    row_number = 0
    
    for row_idx, row in enumerate(rows):
        row_data = [str(cell) if cell is not None else "" for cell in row]
        
        # Saltar filas vac├¡as
        if all(not cell.strip() for cell in row_data):
            continue
        
        # Saltar headers y l├¡neas de encabezado
        first_cell = row_data[0].strip() if row_data else ""
        if any(keyword in first_cell.lower() for keyword in [
            "alejandro", "listado", "art├¡culo", "articulo", "rubro", 
            "existencia", "unidad", "fecha", "hora", "puesto", "usuario",
            "p├íg", "tel:", "fax:", "e-mail", "cba -"
        ]):
            continue
        
        # Buscar celda con formato "CODIGO - NOMBRE"
        producto_cell = ""
        categoria_cell = ""
        stock_cell = ""
        
        for cell in row_data:
            cell_text = cell.strip()
            if not cell_text:
                continue
                
            # Detectar producto (tiene gui├│n y empieza con c├│digo alfanum├®rico)
            # Formato: EX002 - EXHIBIDORA BRIKET 4300
            if re.match(r'^[A-Z0-9]+\s*-\s*.+', cell_text) and not re.search(r'-\s*\d+\s*$', cell_text):
                producto_cell = cell_text
            # Detectar categor├¡a (tiene gui├│n y n├║mero al final)
            # Formato: HELADERAS FREEZER EXHIBIDORAS - 19
            elif re.search(r'-\s*\d+\s*$', cell_text):
                categoria_cell = cell_text
            # Detectar stock (n├║mero decimal)
            elif re.match(r'^\d+[.,]\d+$', cell_text):
                stock_cell = cell_text
        
        # Si encontramos un producto, procesarlo
        if producto_cell:
            row_number += 1
            
            # Extraer c├│digo y nombre
            match = re.match(r'^([A-Z0-9]+)\s*-\s*(.+)$', producto_cell)
            if match:
                codigo = match.group(1).strip()
                nombre = match.group(2).strip()
            else:
                codigo = ""
                nombre = producto_cell
            
            # Log para debug
            logger.info(f"Fila {row_number}: C├│digo='{codigo}', Nombre='{nombre}', Stock='{stock_cell}'")
            
            # Extraer categor├¡a (remover el n├║mero al final)
            categoria_raw = re.sub(r'\s*-\s*\d+\s*$', '', categoria_cell).strip()
            
            # Extraer stock
            stock = 1  # Por defecto 1
            if stock_cell:
                try:
                    # Convertir "1.00" o "1,00" a entero
                    stock_float = float(stock_cell.replace(',', '.'))
                    stock = max(1, int(stock_float))  # M├¡nimo 1
                except ValueError:
                    stock = 1
            
            # Validar y crear producto
            validation = _validate_xlsx_row(
                row_number=row_number,
                nombre=nombre,
                categoria_raw=categoria_raw,
                subcategoria="",
                precio_raw="0",  # Precio 0 por defecto
                stock_raw=str(stock),
                marca=_extract_brand_from_name(nombre),  # Extraer marca del nombre
                descripcion="",
            )
            validations.append(validation)
    
    return validations


def _parse_standard_format(rows: list) -> List[ProductImportValidation]:
    """
    Parsea formato est├índar con headers en primera fila.
    """
    validations: List[ProductImportValidation] = []
    
    if len(rows) < 2:
        return [ProductImportValidation(
            row_number=0,
            valid=False,
            errors=["El archivo no contiene productos para importar"]
        )]
    
    # Primera fila = headers
    raw_headers = [str(cell) if cell is not None else "" for cell in rows[0]]
    column_mapping = _detect_column_mapping(raw_headers)
    
    logger.info(f"Headers detectados: {raw_headers}")
    logger.info(f"Mapeo de columnas: {column_mapping}")
    
    # Procesar filas de datos (desde la fila 2)
    for row_idx, row in enumerate(rows[1:], start=2):
        row_data = [cell for cell in row]
        
        # Saltar filas completamente vac├¡as
        if all(cell is None or str(cell).strip() == "" for cell in row_data):
            continue
        
        # Extraer valores seg├║n el mapeo
        nombre = _get_cell_value(row_data, column_mapping.get("nombre"))
        categoria_raw = _get_cell_value(row_data, column_mapping.get("categoria"))
        subcategoria = _get_cell_value(row_data, column_mapping.get("subcategoria"))
        precio_raw = _get_cell_value(row_data, column_mapping.get("precio"))
        stock_raw = _get_cell_value(row_data, column_mapping.get("stock"))
        marca = _get_cell_value(row_data, column_mapping.get("marca"))
        descripcion = _get_cell_value(row_data, column_mapping.get("descripcion"))
        
        # Validar y construir la fila
        validation = _validate_xlsx_row(
            row_number=row_idx - 1,  # 1-based row number for data rows
            nombre=nombre,
            categoria_raw=categoria_raw,
            subcategoria=subcategoria,
            precio_raw=precio_raw,
            stock_raw=stock_raw,
            marca=marca,
            descripcion=descripcion,
        )
        validations.append(validation)
    
    return validations


def _get_cell_value(row: list, col_idx: Optional[int]) -> str:
    """Obtiene el valor de una celda como string, o vac├¡o si no existe."""
    if col_idx is None or col_idx >= len(row):
        return ""
    val = row[col_idx]
    if val is None:
        return ""
    return str(val).strip()


def _parse_category(categoria_raw: str) -> Tuple[str, str]:
    """
    Parsea la categor├¡a raw del Excel a (categoria, subcategoria_default).
    Si no coincide con las categor├¡as v├ílidas, usa 'electrodomesticos' por defecto.
    """
    cat_lower = categoria_raw.lower().strip()
    
    if cat_lower in ("electrodomesticos", "electrodom├®sticos", "electro"):
        return ("electrodomesticos", "General")
    elif cat_lower in ("muebleria", "muebler├¡a", "muebles"):
        return ("muebleria", "General")
    elif cat_lower in ("colchoneria", "colchoner├¡a", "colchones"):
        return ("colchoneria", "General")
    else:
        # Intentar mapeo por contenido
        return map_categoria(categoria_raw)


def _validate_xlsx_row(
    row_number: int,
    nombre: str,
    categoria_raw: str,
    subcategoria: str,
    precio_raw: str,
    stock_raw: str,
    marca: str,
    descripcion: str,
) -> ProductImportValidation:
    """Valida una fila del Excel y retorna un ProductImportValidation."""
    errors: List[str] = []
    
    # Validar nombre (obligatorio)
    if not nombre.strip():
        errors.append("El campo 'nombre' es obligatorio")
    
    # Parsear precio
    precio = 0.0
    if precio_raw.strip():
        try:
            # Manejar formatos como "1.500,00" o "1500.00"
            cleaned = precio_raw.replace("$", "").replace(" ", "").strip()
            if "," in cleaned and "." in cleaned:
                # Formato 1.500,00 ÔåÆ remover puntos de miles, coma ÔåÆ punto decimal
                cleaned = cleaned.replace(".", "").replace(",", ".")
            elif "," in cleaned:
                cleaned = cleaned.replace(",", ".")
            precio = float(cleaned)
            if precio < 0:
                errors.append("El precio no puede ser negativo")
                precio = 0.0
        except ValueError:
            errors.append(f"Precio inv├ílido: '{precio_raw}'")
    
    # Parsear stock
    stock = 0
    if stock_raw.strip():
        try:
            # Manejar formatos como "10,00" o "10.00"
            cleaned = stock_raw.replace(",", ".").strip()
            stock = int(float(cleaned))
            if stock < 0:
                errors.append("El stock no puede ser negativo")
                stock = 0
        except ValueError:
            errors.append(f"Stock inv├ílido: '{stock_raw}'")
    
    # Parsear categor├¡a
    categoria = "electrodomesticos"
    sub = subcategoria.strip() if subcategoria.strip() else "General"
    if categoria_raw.strip():
        categoria, default_sub = _parse_category(categoria_raw)
        if not subcategoria.strip():
            sub = default_sub
    
    # Si hay errores de validaci├│n, retornar inv├ílido
    if errors:
        return ProductImportValidation(
            row_number=row_number,
            valid=False,
            errors=errors,
        )
    
    try:
        product_row = ProductImportRow(
            nombre=nombre.strip(),
            precio=precio,
            stock=stock,
            categoria=categoria,
            subcategoria=sub,
            marca=marca.strip() if marca.strip() else "Sin marca",
            descripcion=descripcion.strip() if descripcion.strip() else "",
            slug=generate_slug(nombre.strip()),
        )
        
        return ProductImportValidation(
            row_number=row_number,
            valid=True,
            data=product_row,
        )
    except Exception as e:
        return ProductImportValidation(
            row_number=row_number,
            valid=False,
            errors=[f"Error de validaci├│n: {str(e)}"],
        )


async def process_xlsx_import(
    confirmed_rows: List[ExcelImportConfirmRow],
    supabase_client,
) -> BulkImportResponse:
    """
    Procesa la importaci├│n de filas confirmadas desde el preview de Excel.
    
    Args:
        confirmed_rows: Lista de filas confirmadas por el usuario
        supabase_client: Cliente de Supabase para insertar en BD
        
    Returns:
        Resultado de la importaci├│n
    """
    validations: List[ProductImportValidation] = []
    imported_count = 0
    
    for idx, row in enumerate(confirmed_rows, start=1):
        slug = generate_slug(row.nombre)
        
        try:
            product_data = {
                'nombre': row.nombre,
                'slug': slug,
                'categoria': row.categoria,
                'subcategoria': row.subcategoria or 'General',
                'precio': row.precio,
                'precio_original': None,
                'stock': row.stock,
                'marca': row.marca or 'Sin marca',
                'descripcion': row.descripcion or '',
                'imagenes': [row.imagen] if row.imagen else [],
                'especificaciones': {},
                'destacado': False,
                'calificacion': 0.0,
                'cantidad_resenas': 0,
            }
            
            result = supabase_client.table('productos').insert(product_data).execute()
            
            if result.data:
                imported_count += 1
                validations.append(ProductImportValidation(
                    row_number=idx,
                    valid=True,
                    data=ProductImportRow(
                        nombre=row.nombre,
                        precio=row.precio,
                        stock=row.stock,
                        categoria=row.categoria,
                        subcategoria=row.subcategoria or 'General',
                        marca=row.marca or 'Sin marca',
                        descripcion=row.descripcion or '',
                        slug=slug,
                        imagen=row.imagen,
                    ),
                ))
            else:
                validations.append(ProductImportValidation(
                    row_number=idx,
                    valid=False,
                    errors=["Error al insertar en la base de datos"],
                ))
        except Exception as e:
            validations.append(ProductImportValidation(
                row_number=idx,
                valid=False,
                errors=[f"Error al insertar: {str(e)}"],
            ))
    
    valid_count = sum(1 for v in validations if v.valid)
    invalid_count = sum(1 for v in validations if not v.valid)
    
    return BulkImportResponse(
        total_rows=len(confirmed_rows),
        valid_rows=valid_count,
        invalid_rows=invalid_count,
        imported_count=imported_count,
        validations=validations,
        message=f"Importaci├│n completada: {imported_count} productos importados de {len(confirmed_rows)} confirmados",
    )


def map_categoria(categoria_texto: str) -> Tuple[str, str]:
    """
    Mapea la categor├¡a del documento a nuestras categor├¡as de BD.
    
    Args:
        categoria_texto: Texto como "SMART TV - 18"
    
    Returns:
        Tupla (categoria, subcategoria)
    """
    # TODO: Implementar l├│gica de mapeo seg├║n tus reglas de negocio
    # Por ahora, asumimos que todo es electrodom├®sticos
    categoria_lower = categoria_texto.lower()
    
    if 'tv' in categoria_lower or 'smart' in categoria_lower:
        return ('electrodomesticos', 'Televisores')
    elif 'heladera' in categoria_lower or 'freezer' in categoria_lower:
        return ('electrodomesticos', 'Heladeras')
    elif 'lavarropas' in categoria_lower:
        return ('electrodomesticos', 'Lavarropas')
    elif 'colchon' in categoria_lower or 'sommier' in categoria_lower:
        return ('colchoneria', 'Colchones')
    elif 'mueble' in categoria_lower or 'mesa' in categoria_lower or 'silla' in categoria_lower:
        return ('muebleria', 'Muebles')
    else:
        # Por defecto
        return ('electrodomesticos', 'General')


def parse_doc_file(file_content: bytes) -> List[ProductImportValidation]:
    """
    Parsea un archivo .doc/.docx y extrae los productos.
    
    Formato esperado:
    - L├¡nea: "CODIGO - NOMBRE DEL PRODUCTO"
    - Siguiente l├¡nea: "CATEGORIA - XX"
    - Stock: n├║mero en la columna derecha
    
    Args:
        file_content: Contenido del archivo .doc en bytes
    
    Returns:
        Lista de validaciones de productos
    """
    validations: List[ProductImportValidation] = []
    
    try:
        # Intentar primero como .docx (formato ZIP)
        try:
            doc = Document(BytesIO(file_content))
            paragraphs = _collect_docx_lines(doc)
        except zipfile.BadZipFile:
            # Si falla, es un .doc antiguo - extraer texto plano
            text_content = file_content.decode('latin-1', errors='ignore')
            text_content = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text_content)
            paragraphs = [line.strip() for line in text_content.split('\n') if line.strip()]
        
        # DEBUG: Imprimir las primeras 30 l├¡neas para ver qu├® estamos leyendo
        logger.info("=== DEBUG: Primeras 30 l├¡neas del documento ===")
        for idx, para in enumerate(paragraphs[:30]):
            logger.info(f"L├¡nea {idx}: '{para}'")
        logger.info("=== FIN DEBUG ===")
        logger.info(f"Total de l├¡neas en el documento: {len(paragraphs)}")
        
        row_number = 0
        i = 0
        
        while i < len(paragraphs):
            text = paragraphs[i].strip()
            
            # Ignorar encabezados y l├¡neas vac├¡as
            if (not text or 
                'Alejandro Accietto' in text or
                'Listados de Existencias' in text or
                'Rubro' in text or
                'Art├¡culo' in text or
                'Existencia' in text or
                'Unidad Venta' in text or
                'CBA - SANTY HOGAR' in text or
                'P├íg.' in text or
                'Puesto:' in text or
                'Usuario:' in text or
                'Fecha' in text or
                'Hora' in text or
                re.match(r'^\d{2}/\d{2}/\d{4}$', text) or  # Fechas
                re.match(r'^\d{2}:\d{2}$', text)):  # Horas
                i += 1
                continue

            # Formato: solo c├│digo (sin " - "), nombre en la l├¡nea siguiente
            solo_codigo = re.match(r'^([A-Za-z0-9][A-Za-z0-9._-]{2,})\s*$', text)
            if solo_codigo and i + 1 < len(paragraphs):
                nombre_cand = paragraphs[i + 1].strip()
                looks_cat = bool(re.match(r'^.+\s*-\s*\d+$', nombre_cand))
                looks_stock = bool(re.match(r'^(\d+)[,.](\d+)$', nombre_cand))
                if nombre_cand and not looks_cat and not looks_stock:
                    row_number += 1
                    codigo = solo_codigo.group(1).strip()
                    nombre = nombre_cand
                    categoria_texto = ""
                    stock = 0
                    j = i + 2
                    if j < len(paragraphs):
                        nl = paragraphs[j].strip()
                        if re.match(r'^.+\s*-\s*\d+$', nl):
                            categoria_texto = re.sub(
                                r"\s*-\s*\d+$", "", nl
                            ).strip()
                            j += 1
                    if j < len(paragraphs):
                        sl = paragraphs[j].strip()
                        sm = re.match(r"^(\d+)[,.](\d+)$", sl)
                        if sm:
                            stock = int(sm.group(1))
                            j += 1
                    validations.append(
                        _validate_product(
                            row_number,
                            {
                                "codigo": codigo,
                                "nombre": nombre,
                                "categoria": categoria_texto,
                                "stock": stock,
                            },
                        )
                    )
                    i = j
                    continue
            
            # Detectar l├¡nea de producto:
            # - Formato A: "CODIGO - NOMBRE"
            # - Formato B: "CODIGO -" (o solo CODIGO) y el nombre viene en la l├¡nea siguiente
            producto_match = re.match(r'^([A-Z0-9]+)\s*-\s*(.*)$', text)
            
            if producto_match:
                row_number += 1
                codigo = producto_match.group(1).strip()
                nombre = (producto_match.group(2) or "").strip()
                
                logger.info(f"DEBUG: Producto encontrado - C├│digo: {codigo}, Nombre inicial: {nombre}")

                # Si no vino nombre en la misma l├¡nea, tomar la siguiente como nombre/descripcion.
                if not nombre and i + 1 < len(paragraphs):
                    candidate = paragraphs[i + 1].strip()
                    # Evitar tomar encabezados u otras l├¡neas t├¡picas.
                    if candidate and not re.match(r'^.+\s*-\s*\d+$', candidate):
                        nombre = candidate
                        logger.info(f"DEBUG: Nombre tomado de l├¡nea siguiente: {nombre}")
                        i += 1  # Consumimos la l├¡nea del nombre
                
                # Buscar stock en la misma l├¡nea (n├║meros al final)
                stock = 0
                stock_match = re.search(r'(\d+)[,.](\d+)\s*$', text)
                if stock_match:
                    # Remover el stock del nombre
                    nombre = re.sub(r'\s*\d+[,.]\d+\s*$', '', nombre).strip()
                    stock = int(stock_match.group(1))
                    logger.info(f"DEBUG: Stock encontrado en misma l├¡nea: {stock}")
                
                # La siguiente l├¡nea deber├¡a ser la categor├¡a
                categoria_texto = ''
                if i + 1 < len(paragraphs):
                    next_line = paragraphs[i + 1].strip()
                    logger.info(f"DEBUG: Siguiente l├¡nea: '{next_line}'")
                    # Verificar si es una categor├¡a (formato: "TEXTO - XX")
                    if re.match(r'^.+\s*-\s*\d+$', next_line):
                        categoria_texto = re.sub(r'\s*-\s*\d+$', '', next_line).strip()
                        logger.info(f"DEBUG: Categor├¡a encontrada: {categoria_texto}")
                        i += 1  # Saltar la l├¡nea de categor├¡a
                
                # Si no encontramos stock en la l├¡nea del producto, buscar en la siguiente
                if stock == 0 and i + 1 < len(paragraphs):
                    next_line = paragraphs[i + 1].strip()
                    stock_match = re.match(r'^(\d+)[,.](\d+)$', next_line)
                    if stock_match:
                        stock = int(stock_match.group(1))
                        logger.info(f"DEBUG: Stock encontrado en l├¡nea siguiente: {stock}")
                        i += 1  # Saltar la l├¡nea de stock
                
                # Crear el producto
                product_data = {
                    'codigo': codigo,
                    'nombre': nombre,
                    'categoria': categoria_texto,
                    'stock': stock
                }
                
                logger.info(f"DEBUG: Producto completo: {product_data}")
                validations.append(_validate_product(row_number, product_data))
            
            i += 1
        
    except Exception as e:
        import traceback
        logger.error(f"ERROR: {str(e)}")
        logger.error(traceback.format_exc())
        validations.append(
            ProductImportValidation(
                row_number=0,
                valid=False,
                errors=[f"Error al leer el archivo: {str(e)}"]
            )
        )
    
    return validations


def _validate_product(row_number: int, data: dict) -> ProductImportValidation:
    """Valida y crea un ProductImportRow desde los datos parseados."""
    errors = []
    
    # Validar campos requeridos
    if not (data.get('nombre') or '').strip():
        errors.append("Falta el nombre del producto")
    
    # Si hay errores, retornar inv├ílido
    if errors:
        return ProductImportValidation(
            row_number=row_number,
            valid=False,
            errors=errors
        )
    
    try:
        # Mapear categor├¡a
        categoria, subcategoria = map_categoria(data.get('categoria', ''))
        
        # Extraer marca del nombre del producto
        nombre = data['nombre']
        marca = _extract_brand_from_name(nombre)

        product_row = ProductImportRow(
            nombre=nombre,
            precio=0.0,
            stock=data.get('stock', 0),
            categoria=categoria,
            subcategoria=subcategoria,
            marca=marca,
            slug=generate_slug(nombre)
        )
        
        return ProductImportValidation(
            row_number=row_number,
            valid=True,
            data=product_row
        )
    
    except Exception as e:
        return ProductImportValidation(
            row_number=row_number,
            valid=False,
            errors=[f"Error de validaci├│n: {str(e)}"]
        )


async def process_bulk_import(
    file_content: bytes,
    supabase_client,
) -> BulkImportResponse:
    """
    Procesa la importaci├│n masiva de productos.
    
    Args:
        file_content: Contenido del archivo .doc
        supabase_client: Cliente de Supabase para insertar en BD
    
    Returns:
        Resultado de la importaci├│n
    """
    # Parsear el archivo
    validations = parse_doc_file(file_content)
    
    # Contar v├ílidos e inv├ílidos
    valid_rows = [v for v in validations if v.valid]
    invalid_rows = [v for v in validations if not v.valid]
    
    imported_count = 0
    
    # Insertar productos v├ílidos
    for validation in valid_rows:
        if validation.data:
            try:
                # Preparar datos para inserci├│n
                product_data = {
                    'nombre': validation.data.nombre,
                    'slug': validation.data.slug,
                    'categoria': validation.data.categoria,
                    'subcategoria': validation.data.subcategoria,
                    'precio': validation.data.precio,
                    'precio_original': validation.data.precio_costo,
                    'stock': validation.data.stock,
                    'marca': validation.data.marca,
                    'descripcion': validation.data.descripcion or '',
                    'imagenes': [],  # Sin im├ígenes por defecto
                    'especificaciones': {},
                    'destacado': False,
                    'calificacion': 0.0,
                    'cantidad_resenas': 0,
                }
                
                # Insertar en Supabase
                result = supabase_client.table('productos').insert(product_data).execute()
                
                if result.data:
                    imported_count += 1
            
            except Exception as e:
                validation.valid = False
                validation.errors.append(f"Error al insertar: {str(e)}")
    
    return BulkImportResponse(
        total_rows=len(validations),
        valid_rows=len(valid_rows),
        invalid_rows=len(invalid_rows),
        imported_count=imported_count,
        validations=validations,
        message=f"Importaci├│n completada: {imported_count} productos importados de {len(valid_rows)} v├ílidos"
    )
